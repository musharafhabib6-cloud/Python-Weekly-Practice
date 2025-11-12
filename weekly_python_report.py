from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# ===== TITLE PAGE =====
title = doc.add_heading("Weekly Python Practice Report", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nName: ___________________________", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Week: ___________________________", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("\n", style="Normal")

# ===== SECTION 1 =====
doc.add_heading("1. Debugging & Configuration", level=2)
doc.add_paragraph(
    "Task: Created a small Python program with a logical bug and used the debugger to find and fix the issue. "
    "Learned how to use breakpoints, inspect variables, and understand code execution flow."
)

doc.add_paragraph("Example code with a logical bug:")
doc.add_paragraph(
    """\
# Example of a logical bug
def add_numbers(a, b):
    return a - b  # Bug: should be '+' instead of '-'

result = add_numbers(10, 5)
print("The sum is:", result)
""",
    style="List Bullet"
)

doc.add_paragraph(
    "By debugging step-by-step, I found that the subtraction operator caused the wrong output. "
    "After fixing it to '+', the program worked correctly. This improved my understanding of debugging tools."
)

# ===== SECTION 2 =====
doc.add_heading("2. Python Basics", level=2)
doc.add_paragraph(
    "Task: Practiced variables, loops, functions, and lists. Created small programs like a calculator and grade tracker."
)

doc.add_paragraph("Example calculator code:")
doc.add_paragraph(
    """\
def calculator(a, b, operation):
    if operation == 'add':
        return a + b
    elif operation == 'subtract':
        return a - b
    elif operation == 'multiply':
        return a * b
    elif operation == 'divide':
        return a / b

print(calculator(10, 5, 'add'))
""",
    style="List Bullet"
)

doc.add_paragraph(
    "Learned how functions make code reusable and how loops simplify repetitive tasks."
)

# ===== SECTION 3 =====
doc.add_heading("3. NumPy & Pandas", level=2)
doc.add_paragraph(
    "Task: Loaded and processed a CSV file (grades.csv) using Pandas. Calculated averages, sorted data, and filtered results."
)

doc.add_paragraph("Example code:")
doc.add_paragraph(
    """\
import pandas as pd

data = pd.read_csv('grades.csv')
print(data.head())

average = data['Grade'].mean()
print("Average grade:", average)

filtered = data[data['Grade'] > 80]
print(filtered)
""",
    style="List Bullet"
)

doc.add_paragraph(
    "Learned how Pandas simplifies data handling and analysis tasks with simple, powerful functions."
)

# ===== SECTION 4 =====
doc.add_heading("4. Practical OOP (Functions & Classes)", level=2)
doc.add_paragraph(
    "Task: Refactored Python code into a class. Created a 'Student' class to manage data and calculate averages."
)

doc.add_paragraph("Example code:")
doc.add_paragraph(
    """\
class Student:
    def __init__(self, name, marks):
        self.name = name
        self.marks = marks

    def calculate_average(self):
        return sum(self.marks) / len(self.marks)

student1 = Student("Ali", [85, 90, 78])
print(student1.name, "average:", student1.calculate_average())
""",
    style="List Bullet"
)

doc.add_paragraph(
    "Learned how classes improve code structure, making it easier to manage data and add more features later."
)

# ===== SECTION 5 =====
doc.add_heading("5. Mini Project Checkpoint", level=2)
doc.add_paragraph(
    "Task: Combined previous concepts to analyze grades from a CSV file using an object-oriented approach. "
    "Calculated averages and exported results to output.csv."
)

doc.add_paragraph("Example project code:")
doc.add_paragraph(
    """\
import pandas as pd

class GradeAnalyzer:
    def __init__(self, file_path):
        self.data = pd.read_csv(file_path)

    def calculate_average(self):
        return self.data['Grade'].mean()

    def save_results(self, output_path):
        avg = self.calculate_average()
        result = pd.DataFrame({'Average Grade': [avg]})
        result.to_csv(output_path, index=False)

analyzer = GradeAnalyzer('grades.csv')
print("Average grade:", analyzer.calculate_average())
analyzer.save_results('output.csv')
""",
    style="List Bullet"
)

doc.add_paragraph(
    "This project helped strengthen my understanding of combining data analysis with object-oriented programming. "
    "It was my first self-contained Python project producing meaningful output."
)

# ===== SUMMARY =====
doc.add_heading("Summary", level=2)
doc.add_paragraph(
    "This week’s practice included debugging, Python basics, data analysis with NumPy and Pandas, and object-oriented programming. "
    "I learned how to organize projects, write cleaner code, and manage real-world data efficiently."
)

# Save the file
doc.save("Weekly_Python_Practice_Report_Detailed.docx")

print("✅ Report created successfully: Weekly_Python_Practice_Report_Detailed.docx")
